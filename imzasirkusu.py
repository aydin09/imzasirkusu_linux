from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
from docx.shared import Cm, Inches
from docx.shared import Length
import os
from tkinter import *
import sqlite3
import tkinter.ttk as ttk

def bilgi_girişi(event):
    liste1=liste.get(ACTIVE)

    personel_adi_soyadi.delete(0,END)
    personel_gorev_brans.delete(0,END)
                       
    vt = sqlite3.connect(str(liste1)+'.sq3')
    im= vt.cursor()
    im.execute(""" SELECT * FROM imza""")
    rows = im.fetchall()
    data_str = ""
    sf = "{}{}"
    for row in rows:
        data_str += sf.format(row[0], row[1])

        personel_adi_soyadi.insert(END,row[0])
        personel_gorev_brans.insert(END,row[1])
                        
def kaydet():
    kaymakamlik1 = kaymakamlik.get()
    okul_adi1 = okul_adi.get()
    personel_adi_soyadi1 = personel_adi_soyadi.get()
    personel_gorev_brans1 = personel_gorev_brans.get()
    
    
    personel_adi_soyadi.delete(0,END)
    personel_gorev_brans.delete(0,END)

    vt1 = sqlite3.connect('kaymakamlik.sq')
    im1= vt1.cursor()
    im1.execute("""CREATE TABLE IF NOT EXISTS imza(kaymakamlik TEXT)""")
    im1.execute("""UPDATE imza SET  kaymakamlik=?""",(kaymakamlik1,))
    vt1.commit()        
    
    vt2 = sqlite3.connect('okuladi.sql')
    im2= vt2.cursor()
    im2.execute("""CREATE TABLE IF NOT EXISTS imza(okuladi TEXT)""")
    im2.execute("""UPDATE imza SET  okuladi=?""",(okul_adi1,))
    vt2.commit()
    
    if os.path.exists(personel_adi_soyadi1+'.sq3')== False:
        vt1 = sqlite3.connect(personel_adi_soyadi1+'.sq3')
        im1= vt1.cursor()
        im1.execute("""CREATE TABLE IF NOT EXISTS imza(personeladisoyadi TEXT, personelgorevbrans TEXT)""")
        im1.execute("""INSERT INTO imza VALUES  (?,?)""",(personel_adi_soyadi1, personel_gorev_brans1,))
        vt1.commit()

        liste.delete(0,END)

        for i in sorted(os.listdir(os.getcwd())):
            if i.endswith('.sq3'):
                liste.insert(END,i[0:-4])

    else:
        vt2 = sqlite3.connect(personel_adi_soyadi1+'.sq3')
        im2= vt2.cursor()
        im2.execute("""CREATE TABLE IF NOT EXISTS imza(personeladisoyadi TEXT, personelgorevbrans TEXT)""")
        im2.execute("""UPDATE imza SET  personeladisoyadi=?, personelgorevbrans=?""",(personel_adi_soyadi1,personel_gorev_brans1,))
        
        vt2.commit()
       
def cikti():
    vt1 = sqlite3.connect('kaymakamlik.sq')
    im1= vt1.cursor()
    im1.execute("""CREATE TABLE IF NOT EXISTS imza(kaymakamlik TEXT)""")
    im1.execute("""SELECT * FROM  imza""")
    rows = im1.fetchall()
    data_str = ""
    sf = "{}"
    for rowkay in rows:
        data_str += sf.format(rowkay[0])
    vt1.commit()        
    
    vt2 = sqlite3.connect('okuladi.sql')
    im2= vt2.cursor()
    im2.execute("""CREATE TABLE IF NOT EXISTS imza(okuladi TEXT)""")
    im2.execute("""SELECT * FROM  imza""")
    rows = im2.fetchall()
    data_str1 = ""
    sf = "{}"
    for rowokul in rows:
        data_str1 += sf.format(rowokul[0])
    vt2.commit()

    a=[]
    for i in sorted(os.listdir(os.getcwd())):
        if i.endswith('.sq3'):
            a.append(i[0:-4])

    data3 = []
 
    while True:
        for s in range(0,len(a)):
            
            vt3= sqlite3.connect(a[s]+'.sq3')
            im3= vt3.cursor()
            im3.execute(""" SELECT * FROM imza""")
            rows3 = im3.fetchall()
            
            for row3 in rows3:
                data3.append(row3)

        break
    
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)    

    paragraph = document.add_paragraph()
    paragraph.add_run("T.C.").bold = True
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.alignment = 1
   
    paragraph = document.add_paragraph()
    paragraph.add_run(rowkay[0]).bold = True
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.alignment = 1

    paragraph = document.add_paragraph()
    paragraph.add_run(rowokul[0]).bold = True
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.alignment = 1

    paragraph = document.add_paragraph()
    paragraph.add_run("İMZA SİRKÜSÜ"+"\n").bold = True
    paragraph.alignment = 1
             
    table = document.add_table(rows=len(a)+1, cols=4,style = 'Table Grid')
    
    cell = table.cell(0,0)
    table.cell(0,0).paragraphs[0].add_run("SIRA NO").bold = True
    table.cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(0.6)
    
    
    cell = table.cell(0,1)
    table.cell(0,1).paragraphs[0].add_run("PERSONELİN ADI SOYADI").bold = True
    table.cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[1].width = Inches(2.5)

    cell = table.cell(0,2)
    table.cell(0,2).paragraphs[0].add_run("GÖREVİ/BRANŞI").bold = True
    table.cell(0,2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[2].width = Inches(2.5)

    cell = table.cell(0,3)
    table.cell(0,3).paragraphs[0].add_run("İMZA").bold = True
    table.cell(0,3).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[3].width = Inches(1.0)

    for s in range(0,len(a)):
        cell = table.cell(s+1,1)
        cell.text =data3[s][0]
        table.cell(s+1,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        cell = table.cell(s+1,2)
        cell.text =data3[s][1]
        table.cell(s+1,2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT

        cell = table.cell(s+1,0)
        cell.text =str(int(s)+1)
        table.cell(s+1,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    document.save('imzasirkusu.docx')

    os.system("libreoffice --writer imzasirkusu.docx")

def sil():
    data_sil=liste.get(ACTIVE)

    os.remove(data_sil+".sq3")

    liste.delete(0,END)

    for i in sorted(os.listdir(os.getcwd())):
        if i.endswith('.sq3'):
            liste.insert(END,i[0:-4])

   
    personel_adi_soyadi.delete(0,END)
    personel_gorev_brans.delete(0,END)

root = Tk()
root.title("İmza Sirküsü Programı")
root.resizable(width=FALSE ,height=FALSE)
img=PhotoImage(file='imza.png')
root.tk.call('wm','iconphoto',root._w,img)
mainframe = ttk.Frame(root,padding='3 3 12 12')
mainframe.grid(column=0, row=0)
mainframe.columnconfigure(0, weight=1)
mainframe.rowconfigure(0, weight =1)

kaymakamlik = ttk.Entry(mainframe, width =50)
kaymakamlik.grid(column = 2, row = 0)

okul_adi = ttk.Entry(mainframe, width =50)
okul_adi.grid(column = 2, row = 1)

personel_adi_soyadi = ttk.Entry(mainframe, width =50)
personel_adi_soyadi.grid(column = 2, row = 2)

personel_gorev_brans = ttk.Entry(mainframe, width =50)
personel_gorev_brans.grid(column = 2, row = 3)

ttk.Label(mainframe, text ='KAYMAKAMLIK ADI').grid(column = 1, row = 0)
ttk.Label(mainframe, text ='OKULUN ADI').grid(column = 1, row = 1)
ttk.Label(mainframe, text ='PERSONELİN ADI SOYADI').grid(column = 1, row=2)
ttk.Label(mainframe, text ='PERSONELİN GÖREVİ/BRANŞI').grid(column = 1, row=3)

liste = Listbox(mainframe,width=50)
liste.grid(column=3, row=0,rowspan=30,  sticky=(N,S,E,W))
liste.bind("<Double-Button-1>",bilgi_girişi)

kaydirma = ttk.Scrollbar(mainframe, orient="vertical",command=liste.yview)
kaydirma.grid(column=4, row=0, rowspan=30,sticky='ns')

liste.config(yscrollcommand=kaydirma.set)
kaydirma.config(command=liste.yview)

for i in sorted(os.listdir(os.getcwd())):
    if i.endswith('.sq3'):
        liste.insert(END,i[0:-4])

vt1 = sqlite3.connect('okuladi.sql')
im1= vt1.cursor()
im1.execute("""CREATE TABLE IF NOT EXISTS imza(okuladi TEXT)""")
im1.execute("""SELECT * FROM  imza""")
rows = im1.fetchall()
data_str = ""
sf = "{}"
for row1 in rows:
    data_str += sf.format(row1[0])

okul_adi.insert(END,row1[0])

vt1.commit()

vt1 = sqlite3.connect('kaymakamlik.sq')
im1= vt1.cursor()
im1.execute("""CREATE TABLE IF NOT EXISTS imza(kaymakamlik TEXT)""")
im1.execute("""SELECT * FROM  imza""")
rows = im1.fetchall()
data_str = ""
sf = "{}"
for row in rows:
    data_str += sf.format(row[0])

kaymakamlik.insert(END,row[0])

vt1.commit()  

ttk.Button(mainframe, text='Kaydet/Güncelle',command= kaydet).grid(column=1, row=4)
ttk.Button(mainframe, text='Sil', command= sil).grid(column=1, row=5)
ttk.Button(mainframe, text='LibreOffice Writer Ön İzleme', command = cikti).grid(column=1, row=6)

kaymakamlik.focus()

root.mainloop()    
